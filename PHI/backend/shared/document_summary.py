"""Smart Document Summary and Approval System."""

import os
from typing import Dict
import logging

logger = logging.getLogger(__name__)

try:
    import PyPDF2
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def get_pdf_summary(path: str, preview_words: int = 100) -> Dict:
    """Get summary of PDF without reading full content."""
    if not HAS_PYPDF:
        return {"error": "PyPDF2 not installed"}
    
    try:
        with open(path, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            num_pages = len(pdf.pages)
            
            # Get first page preview
            preview = ""
            if num_pages > 0:
                text = pdf.pages[0].extract_text() or ""
                words = text.split()[:preview_words]
                preview = " ".join(words)
            
            return {
                "file_type": "pdf",
                "file_path": path,
                "file_name": os.path.basename(path),
                "total_pages": num_pages,
                "file_size_kb": os.path.getsize(path) / 1024,
                "preview": preview[:200] + "..." if len(preview) > 200 else preview,
                "created_date": None,
                "modified_date": None,
                "status": "ready_for_approval"
            }
    except FileNotFoundError:
        return {"error": f"File not found: {path}"}
    except Exception as e:
        return {"error": f"Failed to summarize PDF: {str(e)}"}

def get_docx_summary(path: str, preview_words: int = 100) -> Dict:
    """Get summary of DOCX without reading full content."""
    if not HAS_DOCX:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(path)
        props = doc.core_properties
        
        # Get preview from first few paragraphs
        preview_text = ""
        for para in doc.paragraphs[:3]:
            preview_text += para.text + " "
        
        words = preview_text.split()[:preview_words]
        preview = " ".join(words)
        
        return {
            "file_type": "docx",
            "file_path": path,
            "file_name": os.path.basename(path),
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "file_size_kb": os.path.getsize(path) / 1024,
            "preview": preview[:200] + "..." if len(preview) > 200 else preview,
            "title": props.title or "",
            "author": props.author or "",
            "created_date": str(props.created) if props.created else "",
            "modified_date": str(props.modified) if props.modified else "",
            "status": "ready_for_approval"
        }
    except FileNotFoundError:
        return {"error": f"File not found: {path}"}
    except Exception as e:
        return {"error": f"Failed to summarize DOCX: {str(e)}"}

def get_text_file_summary(path: str, preview_lines: int = 5) -> Dict:
    """Get summary of text file without reading full content."""
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()[:preview_lines]
            preview = "".join(lines)
            total_lines = sum(1 for _ in open(path))
        
        return {
            "file_type": "text",
            "file_path": path,
            "file_name": os.path.basename(path),
            "total_lines": total_lines,
            "file_size_kb": os.path.getsize(path) / 1024,
            "preview": preview[:200] + "..." if len(preview) > 200 else preview,
            "status": "ready_for_approval"
        }
    except FileNotFoundError:
        return {"error": f"File not found: {path}"}
    except Exception as e:
        return {"error": f"Failed to summarize text file: {str(e)}"}

def get_file_summary(path: str) -> Dict:
    """Intelligently summarize any file type."""
    if not os.path.exists(path):
        return {"error": f"File not found: {path}"}
    
    _, ext = os.path.splitext(path.lower())
    
    if ext == ".pdf":
        return get_pdf_summary(path)
    elif ext == ".docx":
        return get_docx_summary(path)
    elif ext in [".txt", ".py", ".js", ".json", ".yaml", ".yml", ".sql", ".env", ".sh", ".bat"]:
        return get_text_file_summary(path)
    else:
        return {
            "file_type": "unknown",
            "file_path": path,
            "file_name": os.path.basename(path),
            "file_size_kb": os.path.getsize(path) / 1024,
            "error": f"Unsupported file type: {ext}",
            "status": "unsupported"
        }
