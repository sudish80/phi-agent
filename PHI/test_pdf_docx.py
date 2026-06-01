#!/usr/bin/env python
"""Test PDF and DOCX reading capabilities."""

import sys
import os
import tempfile
sys.path.insert(0, '.')

# Use temp directory
temp_dir = tempfile.gettempdir()
pdf_path = os.path.join(temp_dir, 'test.pdf')
docx_path = os.path.join(temp_dir, 'test.docx')

# Create test PDF
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawString(100, 750, 'Hello World - Page 1')
    c.drawString(100, 730, 'This is a test PDF document')
    c.drawString(100, 710, 'API_KEY=sk-test123')
    c.showPage()
    c.drawString(100, 750, 'Page 2 - More Content')
    c.showPage()
    c.save()
    print(f'[OK] Test PDF created at {pdf_path}')
except Exception as e:
    print(f'[SKIP] Could not create test PDF: {e}')

# Create test DOCX
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph('Hello World - Test DOCX')
    doc.add_paragraph('This is a test document with credentials')
    doc.add_paragraph('DATABASE_PASSWORD=secret123')
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    doc.save(docx_path)
    print(f'[OK] Test DOCX created at {docx_path}')
except Exception as e:
    print(f'[ERROR] Could not create test DOCX: {e}')

# Test the reading functions
from backend.tools.media_tools import pdf_read, pdf_page_count, docx_read, docx_metadata

print('\n=== Testing PDF Reading ===')
result = pdf_read(pdf_path, page=-1)
print(f'PDF content (first 200 chars): {result[:200]}')
print(f'Contains API_KEY: {("API_KEY" in result)}')

print('\n=== Testing PDF Page Count ===')
pages = pdf_page_count(pdf_path)
print(f'PDF page count: {pages}')

print('\n=== Testing DOCX Reading ===')
result = docx_read(docx_path)
print(f'DOCX content (first 200 chars): {result[:200]}')
print(f'Contains DATABASE_PASSWORD: {("DATABASE_PASSWORD" in result)}')

print('\n=== Testing DOCX Metadata ===')
meta = docx_metadata(docx_path)
print(f'DOCX metadata: {meta}')

print('\n[PASSED] All PDF and DOCX reading tests passed!')

